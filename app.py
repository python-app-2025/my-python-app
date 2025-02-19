import streamlit as st
import sqlite3
import pandas as pd
from datetime import datetime, time
from docx import Document
import io
import os
import uuid
import matplotlib.pyplot as plt
from openpyxl import Workbook
from openpyxl.drawing.image import Image as OpenpyxlImage



# Настройки страницы
st.set_page_config(
    page_title="Проверки ОТиПБ",
    page_icon="✅",
    layout="wide"
)


# Общие настройки
DATABASE_NAME = "inspections.db"
COMMON_DB = "common.db"
SOFTWARE_DB = "software_checks.db"
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# --------------------------
# Общие функции для работы с организациями
# --------------------------

def init_common_db():
    conn = sqlite3.connect(COMMON_DB)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS organizations
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  name TEXT UNIQUE)''')
    conn.commit()
    conn.close()

def add_organization(name):
    conn = sqlite3.connect(COMMON_DB)
    try:
        c = conn.cursor()
        c.execute("INSERT INTO organizations (name) VALUES (?)", (name,))
        conn.commit()
    except sqlite3.IntegrityError:
        raise ValueError("Организация с таким названием уже существует")
    finally:
        conn.close()

def get_organizations():
    conn = sqlite3.connect(COMMON_DB)
    c = conn.cursor()
    c.execute("SELECT name FROM organizations ORDER BY name")
    organizations = [row[0] for row in c.fetchall()]
    conn.close()
    return organizations

def delete_organization(name):
    conn = sqlite3.connect(COMMON_DB)
    c = conn.cursor()
    c.execute("DELETE FROM organizations WHERE name=?", (name,))
    conn.commit()
    conn.close()

# Обновление организации
def update_organization(old_name, new_name):
    if not old_name or not new_name:
        raise ValueError("Название организации не может быть пустым.")
    
    if old_name == new_name:
        raise ValueError("Новое название должно отличаться от старого.")
    
    conn = sqlite3.connect(COMMON_DB)
    c = conn.cursor()
    try:
        c.execute("SELECT name FROM organizations WHERE name = ?", (new_name,))
        if c.fetchone():
            raise ValueError(f"Организация с названием '{new_name}' уже существует.")
        
        c.execute("UPDATE organizations SET name = ? WHERE name = ?", (new_name, old_name))
        conn.commit()
    except sqlite3.Error as e:
        raise ValueError(f"Ошибка при обновлении организации: {e}")
    finally:
        conn.close()

        
def get_record_by_id(record_id):
    conn = sqlite3.connect(SOFTWARE_DB)
    c = conn.cursor()
    c.execute("SELECT * FROM checks WHERE id=?", (record_id,))
    record = c.fetchone()
    conn.close()
    return record
        
# --------------------------
# Главное меню
# --------------------------

def main_menu():
    st.sidebar.title("Главное меню")
    if "module" not in st.session_state:
        st.session_state.module = None

    if st.sidebar.button("📋 Проверки ОТиПБ"):
        st.session_state.module = "module1"
    
    if st.sidebar.button("🏗️ Проверки в СП"):
        st.session_state.module = "module2"
    
    if st.sidebar.button("🏢 Список ПО"):
        st.session_state.module = "module3"
    
    if st.sidebar.button("🚪 Выход"):
        st.session_state.module = None
        st.rerun()

# --------------------------
# Модуль 1: Проверки ОТиПБ
# --------------------------

def module1():
    st.title("📋 Управление проверками ОТиПБ")

    # Функции БД
    def create_db():
        conn = sqlite3.connect(DATABASE_NAME)
        c = conn.cursor()
        c.execute('''CREATE TABLE IF NOT EXISTS inspections
                     (id INTEGER PRIMARY KEY AUTOINCREMENT,
                      inspection_date TEXT,
                      object TEXT,
                      section TEXT,
                      organization TEXT,
                      violator_name TEXT,
                      violation_description TEXT,
                      violation_type TEXT,
                      violation_category TEXT,
                      risk_level TEXT,
                      inspector_name TEXT,
                      elimination_date TEXT,
                      elimination_status TEXT,
                      photo_path TEXT)''')
        conn.commit()
        conn.close()

    def add_to_db(data):
        conn = sqlite3.connect(DATABASE_NAME)
        c = conn.cursor()
        c.execute('''INSERT INTO inspections VALUES 
                  (NULL,?,?,?,?,?,?,?,?,?,?,?,?,?)''', data)
        conn.commit()
        conn.close()

    def get_all_data():
        conn = sqlite3.connect(DATABASE_NAME)
        df = pd.read_sql('SELECT * FROM inspections', conn)
        conn.close()
        return df

    def update_db(data):
        conn = sqlite3.connect(DATABASE_NAME)
        c = conn.cursor()
        c.execute('''UPDATE inspections SET
                  inspection_date=?,
                  object=?,
                  section=?,
                  organization=?,
                  violator_name=?,
                  violation_description=?,
                  violation_type=?,
                  violation_category=?,
                  risk_level=?,
                  inspector_name=?,
                  elimination_date=?,
                  elimination_status=?,
                  photo_path=?
                  WHERE id=?''', data)
        conn.commit()
        conn.close()

    def delete_from_db(record_id):
        conn = sqlite3.connect(DATABASE_NAME)
        c = conn.cursor()
        c.execute('SELECT photo_path FROM inspections WHERE id=?', (record_id,))
        result = c.fetchone()
        if result and result[0] and os.path.exists(result[0]):
            os.remove(result[0])
        c.execute('DELETE FROM inspections WHERE id=?', (record_id,))
        conn.commit()
        conn.close()

    # Форма добавления записи
    create_db()
    with st.expander("➕ Добавить новую запись", expanded=True):
        with st.form("add_form", clear_on_submit=True):
            cols = st.columns(2)
            inspection_date = cols[0].date_input("Дата проверки*", datetime.today())
            object_val = cols[1].selectbox(
                "Объект проверки*",
                ["УТЭЦ-2", "АНГЦ-5", "Стан 2000", "КЦ-1", "КЦ-2", "АТУ", "ДЦ-1", "ДЦ-2","ЦХПП","ЦГП","УЖДТ"]
            )
            
            section = cols[0].selectbox("Участок проверки*", ["Участок монтажа м\к","Сварочный участок","Участок установки оборудования","Ось 11-3","Отметка +45.100","Маслоподвал"])
            organization = cols[1].selectbox(
                "Наименование организации*", 
                get_organizations(),
                index=0
            )
            
            cols_v = st.columns(2)
            violator_name = cols_v[0].text_input("ФИО Нарушителя*")
            violation_description = cols_v[1].text_area("Описание нарушения*")

            
            cols2 = st.columns(2)
            violation_type = cols2[0].selectbox(
                "Тип нарушения*", [
                "Работы на высоте", "Огневые работы/Пожарная безопасность", 
                "Грузоподъёмные работы/Работа с ПС", "Электробезопасность", 
                "Работы в газоопасн. местах/замкнутом простр-ве", 
                "Земляные работы", "Документы/Допуски и удостоверения", 
                "Исправность инструментов и приспособлений", 
                "Применение/Исправность СИЗ", 
                "Содержание территории/рабочих мест", 
                "Безопасность дорожного движения"])
            
            
            # Второй выпадающий список
            violation_category = cols2[1].selectbox("Категория нарушения*", ["Применение СИЗ", "Обучение и аттестации", "ППР", "Леса", "Анкерные линии", "Другое"])
            

            cols3 = st.columns(3)
            risk_level = cols3[1].selectbox(
                "Уровень риска*", 
                ["высокий", "средний", "низкий"]
            )
            
            inspector_name = cols_v[0].selectbox(
                "Проверяющий*", 
                ["Супервайзер ИТК Иванов И.И.", "Специалист ОТиПБ Петров П.П.", "Специалист ОТ ПО Сидоров С.С."]
            )
            elimination_date = cols3[2].date_input(
                "Дата устранения*", 
                datetime.today()
            )
            elimination_status = cols3[0].selectbox(
                "Статус устранения*", 
                ["не устранено", "устранено"]
            )
            
            uploaded_photo = st.file_uploader(
                "Загрузить фото нарушения",
                type=['jpg', 'jpeg', 'png'],
                accept_multiple_files=False
            )
            
            if st.form_submit_button("💾 Сохранить запись"):
                photo_path = save_uploaded_file(uploaded_photo)
                data = (
                    inspection_date.strftime("%d.%m.%Y"),
                    object_val,
                    section,
                    organization,
                    violator_name,
                    violation_description,
                    violation_type,
                    violation_category,
                    risk_level,
                    inspector_name,
                    elimination_date.strftime("%d.%m.%Y"),
                    elimination_status,
                    photo_path
                )
                add_to_db(data)
                st.success("Запись успешно сохранена!")
                st.rerun()

    # Таблица данных
    st.subheader("📊 Список проверок")
    df = get_all_data()
    
    if not df.empty:
        edited_df = st.data_editor(
            df,
            column_config={
                "photo_path": st.column_config.ImageColumn(
                    "Фото",
                    help="Загруженные изображения"
                )
            },
            hide_index=True,
            use_container_width=True,
            disabled=df.columns.tolist()
        )
        
        # Управление записями
        cols = st.columns(4)
        selected_id = cols[0].number_input(
            "Введите ID записи", 
            min_value=1,
            max_value=df['id'].max()
        )
        
        if cols[1].button("🗑️ Удалить запись"):
            delete_from_db(selected_id)
            st.success("Запись удалена!")
            st.rerun()
            
        if cols[2].button("📥 Экспорт в Excel"):
            # Создаем новый Excel-файл
            wb = Workbook()
            ws = wb.active
            # Заголовки столбцов
            ws.append(df.columns.tolist())
    
            # Данные
            for row in df.itertuples(index=False):
                ws.append(row)
    
            # Сохраняем файл в буфер
            output = io.BytesIO()
            wb.save(output)
            output.seek(0)

            # Предлагаем пользователю скачать файл
            st.download_button(
                label="⬇️ Скачать Excel",
                data=output,
                file_name='inspections.xlsx',
                mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    

        
        if cols[3].button("📄 Сформировать акт"):
            record = df[df['id'] == selected_id].iloc[0].to_dict()
            doc_buffer = generate_act(record)
            if doc_buffer:
                st.download_button(
                    label="⬇️ Скачать акт",
                    data=doc_buffer,
                    file_name=f"Акт_{record['id']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        # Просмотр фото
        if selected_id:
            record = df[df['id'] == selected_id]
            if not record.empty:
                photo_path = record.iloc[0]['photo_path']
                if photo_path and os.path.exists(photo_path):
                    st.image(photo_path, caption="Прикрепленное фото", width=300)
                else:
                    st.warning("Для этой записи нет прикрепленного фото")
    else:
        st.info("Нет данных для отображения")

# --------------------------
# Модуль 2: Проверки в СП
# --------------------------

def module2():
    st.title("🏗️ Проверки в СП")


    # Функции БД
    def init_db():
        conn = sqlite3.connect(SOFTWARE_DB)
        c = conn.cursor()
        c.execute('''CREATE TABLE IF NOT EXISTS checks
                     (id INTEGER PRIMARY KEY AUTOINCREMENT,
                      date TEXT,
                      sp_name TEXT,
                      responsible TEXT,
                      po_name TEXT,
                      object TEXT,
                      works_count INTEGER,
                      responsibility_zone TEXT,
                      start_time TEXT,
                      end_time TEXT,
                      personnel_count INTEGER,
                      checks_count INTEGER,
                      violations_count INTEGER,
                      violation_type TEXT,
                      kpb_violation TEXT,
                      kpb_detected INTEGER,
                      act_issued INTEGER)''')
        c.execute('''CREATE TABLE IF NOT EXISTS photos
                     (id INTEGER PRIMARY KEY AUTOINCREMENT,
                      record_id INTEGER,
                      file_path TEXT,
                      FOREIGN KEY(record_id) REFERENCES checks(id))''')
        conn.commit()
        conn.close()

    def add_record(data):
        conn = sqlite3.connect(SOFTWARE_DB)
        c = conn.cursor()
        c.execute('''INSERT INTO checks 
                     (date, sp_name, responsible, po_name, object, works_count, responsibility_zone, 
                      start_time, end_time, personnel_count, checks_count, violations_count, 
                      violation_type, kpb_violation, kpb_detected, act_issued) 
                     VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''', data)
        record_id = c.lastrowid
        conn.commit()
        conn.close()
        return record_id

    def save_photos(record_id, uploaded_files):
        if not uploaded_files:
            return
        save_dir = f"uploads/{record_id}"
        os.makedirs(save_dir, exist_ok=True)
        conn = sqlite3.connect(SOFTWARE_DB)
        c = conn.cursor()
        for uploaded_file in uploaded_files:
            file_path = os.path.join(save_dir, uploaded_file.name)
            with open(file_path, 'wb') as f:
                f.write(uploaded_file.getbuffer())
            c.execute("INSERT INTO photos (record_id, file_path) VALUES (?, ?)", 
                     (record_id, file_path))
        conn.commit()
        conn.close()

    def get_photos(record_id):
        conn = sqlite3.connect(SOFTWARE_DB)
        c = conn.cursor()
        c.execute("SELECT file_path FROM photos WHERE record_id=?", (record_id,))
        photos = [row[0] for row in c.fetchall()]
        conn.close()
        return photos

    def delete_record(record_id):
        conn = sqlite3.connect(SOFTWARE_DB)
        c = conn.cursor()
        photos = get_photos(record_id)
        for photo in photos:
            if os.path.exists(photo):
                os.remove(photo)
        c.execute("DELETE FROM photos WHERE record_id=?", (record_id,))
        c.execute("DELETE FROM checks WHERE id=?", (record_id,))
        conn.commit()
        conn.close()
        dir_path = f"uploads/{record_id}"
        if os.path.exists(dir_path):
            try:
                os.rmdir(dir_path)
            except OSError:
                pass

    def get_records():
        conn = sqlite3.connect(SOFTWARE_DB)
        c = conn.cursor()
        c.execute("SELECT * FROM checks")
        records = c.fetchall()
        conn.close()
        return records

    def update_record(data):
        conn = sqlite3.connect(SOFTWARE_DB)
        c = conn.cursor()
        try:
            c.execute('''UPDATE checks SET
                 date=?,
                 sp_name=?,
                 responsible=?,
                 po_name=?,
                 object=?,
                 works_count=?,
                 responsibility_zone=?,
                 start_time=?,
                 end_time=?,
                 personnel_count=?,
                 checks_count=?,
                 violations_count=?,
                 violation_type=?,
                 kpb_violation=?,
                 kpb_detected=?,
                 act_issued=?
                 WHERE id=?''', data)
            conn.commit()
        except sqlite3.Error as e:
            st.error(f"Ошибка при обновлении записи: {e}")
        finally:
            conn.close()
        
    # Новая функция для получения данных с путями к фото
    def get_all_data():
        conn = sqlite3.connect(SOFTWARE_DB)
        df = pd.read_sql('''
            SELECT c.*, GROUP_CONCAT(p.file_path) as photo_paths 
            FROM checks c
            LEFT JOIN photos p ON c.id = p.record_id
            GROUP BY c.id
        ''', conn)
        conn.close()
        df['photo_paths'] = df['photo_paths'].apply(lambda x: x.split(',') if x else [])
        return df

    
    # Интерфейс модуля
    init_db()
    

    # Форма добавления записи
    with st.expander("➕ Добавить новую запись", expanded=False):
        with st.form("add_record_form", clear_on_submit=True):
            cols = st.columns(2)
            date = cols[0].date_input("Дата*", datetime.today())
            date_str = date.strftime("%d.%m.%Y")  # Форматируем дату
            sp_name = cols[1].selectbox("Наименование СП*", ["АТУ", "ДЦ-1", "ДЦ-2", "КЦ-1","КЦ-2","ЦХПП","ЦГП","УЖДТ"])


            cols1 = st.columns(3)
            responsible = cols1[0].selectbox("Ответственный от СП*", ["Мастер Иванов И.И.", "Начальник участка Петров П.П.", "Главный специалист Сидоров С.С."])
            object = cols1[1].selectbox("Объект/Участок", ["Участок-1", "Участок-2", "Участок-3", "Участок-4"])
            responsibility_zone = cols1[2].selectbox("Зона ответственности (СП)",["АТУ", "ДЦ-1", "ДЦ-2", "КЦ-1","КЦ-2","ЦХПП","ЦГП","УЖДТ"])

            cols2 = st.columns(3)
            po_name = cols2[0].selectbox("Наименование ПО*", get_organizations())          
            start_time = cols2[1].time_input("Время начала работ*", time(8, 0))
            end_time = cols2[2].time_input("Время окончания работ*", time(17, 0))
            
            cols3 = st.columns(2)
            personnel_count = cols3[0].number_input("Кол-во персонала ПО*", min_value=1)
            works_count = cols3[1].number_input("Кол-во выполняемых работ*", min_value=1)

            cols4 = st.columns(2)
            checks_count = cols4[0].number_input("Проведено проверок*", min_value=1)
            violations_count = cols4[1].number_input("Количество нарушений*", min_value=0)
            
            violation_type = st.selectbox("Тип нарушения*", [
                "Работы на высоте", 
                "Огневые работы/Пожарная безопасность", 
                "Грузоподъёмные работы/Работа с ПС", 
                "Электробезопасность", 
                "Работы в газоопасн. местах/замкнутом простр-ве", 
                "Земляные работы", 
                "Документы/Допуски и удостоверения", 
                "Исправность инструментов и приспособлений", 
                "Применение/Исправность СИЗ", 
                "Содержание территории/рабочих мест", 
                "Безопасность дорожного движения", "Нарушений не выявлено"
            ])


            cols5 = st.columns(2)
            kpb_violation = cols5[0].selectbox("Нарушения КПБ*", ["Нет",
                "Нет алкоголю и наркотикам", 
                "Сообщай о происшествиях", 
                "Получи допуск", 
                "Защити себя от падения"
            ])
            
            act_issued = cols5[1].selectbox("Оформлен Акт*", ["Нет", "Да"])
            
            uploaded_files = st.file_uploader(
                "Прикрепить фотографии",
                type=["png", "jpg", "jpeg"],
                accept_multiple_files=True
            )
            
            if st.form_submit_button("💾 Сохранить запись"):
                data = (
                    date_str,  # Используем отформатированную дату
                    sp_name,
                    responsible,
                    po_name,
                    object,
                    works_count,
                    responsibility_zone,
                    start_time.strftime("%H:%M"),
                    end_time.strftime("%H:%M"),
                    personnel_count,
                    checks_count,
                    violations_count,
                    violation_type,
                    kpb_violation,
                    1 if kpb_violation in ("Нет алкоголю и наркотикам", "Сообщай о происшествиях", "Защити себя от падения", "Получи допуск") else 0,
                    1 if act_issued == "Да" else 0
                )
                record_id = add_record(data)
                save_photos(record_id, uploaded_files)
                st.success("Запись успешно сохранена!")
                st.rerun()


    # Отображение данных
    with st.expander("📋 Все записи"):
     records = get_records()
     df = pd.DataFrame(records, columns=[
        "ID", "Дата", "СП", "Ответственный", "ПО", "Объект", 
        "Кол-во работ", "Зона ответ.", "Начало", "Окончание", 
        "Персонал", "Проверки", "Нарушения", "Тип нарушения", 
        "КПБ нарушение", "КПБ выявлено", "Акт"])


# Преобразуем даты в формат дд.мм.гггг
     df["Дата"] = pd.to_datetime(df["Дата"], format="%d.%m.%Y").dt.strftime("%d.%m.%Y")

    
     st.dataframe(
        df.drop(columns=["КПБ выявлено"]),
        use_container_width=True,
        hide_index=True)

# Управление записями
     cols = st.columns(4)
     selected_id = cols[0].number_input("Введите ID записи", min_value=1)

     if cols[1].button("🗑️ Удалить запись"):
        delete_record(selected_id)
        st.success("Запись удалена!")
        st.rerun()

        
     if selected_id:
        photos = get_photos(selected_id)
        if photos:
            cols = st.columns(3)
            for i, photo in enumerate(photos):
                with cols[i % 3]:
                    st.image(photo, use_container_width=True, width=300)
        else:
            st.warning("Нет фото для этой записи")

          
     if st.button("✏️ Редактирование записи"):
        record = get_record_by_id(selected_id)  # Получаем данные записи по ID
        if record:
                with st.form("edit_form"):
                    st.write("Редактирование записи ID:", selected_id)
                    # Предзаполняем поля формы текущими данными
                    edit_date = st.date_input("Дата", datetime.strptime(record[1], "%d.%m.%Y"))
                    edit_sp_name = st.selectbox("СП", ["АТУ", "ДЦ-1", "ДЦ-2", "КЦ-1", "КЦ-2", "ЦХПП", "ЦГП", "УЖДТ"], index=["АТУ", "ДЦ-1", "ДЦ-2", "КЦ-1", "КЦ-2", "ЦХПП", "ЦГП", "УЖДТ"].index(record[2]))
                    edit_responsible = st.text_input("Ответственный", value=record[3])
                    edit_po_name = st.selectbox("ПО", get_organizations(), index=get_organizations().index(record[4]))
                    edit_object = st.text_input("Объект", value=record[5])
                    edit_works_count = st.number_input("Кол-во работ", value=record[6])
                    edit_responsibility_zone = st.text_input("Зона ответственности", value=record[7])
                    edit_start_time = st.time_input("Начало работ", value=datetime.strptime(record[8], "%H:%M").time())
                    edit_end_time = st.time_input("Окончание работ", value=datetime.strptime(record[9], "%H:%M").time())
                    edit_personnel_count = st.number_input("Кол-во персонала", value=record[10])
                    edit_checks_count = st.number_input("Проведено проверок", value=record[11])
                    edit_violations_count = st.number_input("Количество нарушений", value=record[12])
                    edit_violation_type = st.selectbox("Тип нарушения", [
                        "Работы на высоте", "Огневые работы/Пожарная безопасность", 
                        "Грузоподъёмные работы/Работа с ПС", "Электробезопасность", 
                        "Работы в газоопасн. местах/замкнутом простр-ве", 
                        "Земляные работы", "Документы/Допуски и удостоверения", 
                        "Исправность инструментов и приспособлений", 
                        "Применение/Исправность СИЗ", 
                        "Содержание территории/рабочих мест", 
                        "Безопасность дорожного движения", "Нарушений не выявлено"
                    ], index=[
                        "Работы на высоте", "Огневые работы/Пожарная безопасность", 
                        "Грузоподъёмные работы/Работа с ПС", "Электробезопасность", 
                        "Работы в газоопасн. местах/замкнутом простр-ве", 
                        "Земляные работы", "Документы/Допуски и удостоверения", 
                        "Исправность инструментов и приспособлений", 
                        "Применение/Исправность СИЗ", 
                        "Содержание территории/рабочих мест", 
                        "Безопасность дорожного движения", "Нарушений не выявлено"
                    ].index(record[13]))
                    edit_kpb_violation = st.selectbox("Нарушения КПБ", ["Нет", "Нет алкоголю и наркотикам", "Сообщай о происшествиях", "Получи допуск", "Защити себя от падения"], index=["Нет", "Нет алкоголю и наркотикам", "Сообщай о происшествиях", "Получи допуск", "Защити себя от падения"].index(record[14]))
                    edit_act_issued = st.selectbox("Акт оформлен", ["Нет", "Да"], index=0 if record[16] == 0 else 1)

                    
                    edit_date_str = edit_date.strftime("%d.%m.%Y")  # Форматируем дату

                    if st.form_submit_button("Сохранить изменения"):
                        update_data = (
                            edit_date.strftime("%d.%m.%Y"),
                            edit_sp_name,
                            edit_responsible,
                            edit_po_name,
                            edit_object,
                            edit_works_count,
                            edit_responsibility_zone,
                            edit_start_time.strftime("%H:%M"),
                            edit_end_time.strftime("%H:%M"),
                            edit_personnel_count,
                            edit_checks_count,
                            edit_violations_count,
                            edit_violation_type,
                            edit_kpb_violation,
                            1 if kpb_violation in ("Нет алкоголю и наркотикам", "Сообщай о происшествиях", "Защити себя от падения", "Получи допуск") else 0,
                            1 if edit_act_issued =="Да" else 0,
                            selected_id
                        )

                            # Выводим данные для отладки
                        st.write("Данные для обновления:", update_data)
                        try:
                            update_record(update_data)
                            st.success("Изменения сохранены!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Ошибка: {str(e)}")
        else:
             st.error("Запись не найдена.")
                
     if st.button("📥 Экспорт в Excel"):

            # Создаем новый Excel-файл
            wb = Workbook()
            ws = wb.active
            # Заголовки столбцов
            ws.append(df.columns.tolist())
    
            # Данные
            for row in df.itertuples(index=False):
                ws.append(row)
    
            # Сохраняем файл в буфер
            output = io.BytesIO()
            wb.save(output)
            output.seek(0)

            # Предлагаем пользователю скачать файл
            st.download_button(
                label="⬇️ Скачать Excel",
                data=output,
                file_name='sp_checks.xlsx',
                mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

    # Аналитика
    with st.expander("📈 Аналитика и отчеты"):
        po_list = get_organizations()
        selected_po = st.selectbox("Выберите ПО", po_list)
        
        cols = st.columns(2)
        start_date = cols[0].date_input("Начальная дата", datetime.today())
        end_date = cols[1].date_input("Конечная дата", datetime.today())

        # Преобразуем даты в формат дд.мм.гггг
        start_date_str = start_date.strftime("%d.%m.%Y")
        end_date_str = end_date.strftime("%d.%m.%Y")
        
        if st.button("Сгенерировать отчет"):
            conn = sqlite3.connect(SOFTWARE_DB)
            query = f"""
                SELECT date, violations_count 
                FROM checks 
                WHERE po_name = '{selected_po}' 
                AND date BETWEEN '{start_date_str}' 
                AND '{end_date_str}'
            """
            df = pd.read_sql(query, conn)
            conn.close()

            # Преобразуем даты в формат дд.мм.гггг
            df['date'] = pd.to_datetime(df['date'], format="%d.%m.%Y").dt.strftime("%d.%m.%Y")
            
            fig, ax = plt.subplots(figsize=(10, 6))
            ax.plot(df['date'], df['violations_count'], marker='o', linestyle='-') 
            ax.set_xlabel("Дата")
            ax.set_ylabel("Количество нарушений")
            ax.set_title(f"Динамика нарушений для {selected_po}")
            ax.grid(True)
            st.pyplot(fig)

            # Выгрузка в excel
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='Данные', index=False)
                
                buf = io.BytesIO()
                fig.savefig(buf, format='png', bbox_inches='tight')
                buf.seek(0)
                
                workbook = writer.book
                worksheet = workbook.create_sheet('График')
                img = OpenpyxlImage(buf)
                worksheet.add_image(img, 'A1')
            
            output.seek(0)
            st.download_button(
                label="📥 Скачать отчет",
                data=output,
                file_name=f"Отчет_{selected_po}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

# --------------------------
# Модуль 3: Управление организациями
# --------------------------

def module3():
    st.title("🏢 Список ПО")
    init_common_db()

    with st.expander("➕ Добавить организацию", expanded=True):
        with st.form("add_org_form", clear_on_submit=True):
            new_org = st.text_input("Название организации*")
            if st.form_submit_button("Добавить"):
                if new_org:
                    try:
                        add_organization(new_org)
                        st.success("Организация добавлена!")
                        st.rerun()
                    except ValueError as e:
                        st.error(str(e))

    # Вывод списка всех организаций
    with st.expander("📋 Список всех организаций", expanded=False):
     orgs = get_organizations()
     if orgs:
        st.table(orgs)
     else:
        st.warning("Нет зарегистрированных организаций")

    # Пагинация


 # Вывод списка организаций с возможностью редактирования
    with st.expander("✏️ Редактировать список организаций", expanded=False):
     items_per_page = 10  # Количество организаций на странице
     page_number = st.number_input("📄 Номер страницы", min_value=1)
     start_idx = (page_number - 1) * items_per_page
     end_idx = start_idx + items_per_page
     paginated_orgs = orgs[start_idx:end_idx]
     if paginated_orgs:
        for idx, org in enumerate(paginated_orgs):
            cols = st.columns([3, 1, 1])
            cols[0].write(org)  # Отображение названия организации
            edit_org = cols[1].text_input(f"Редактировать {org}", key=f"edit_{org}", value=org)
            if cols[2].button("✏️ Сохранить", key=f"save_{org}"):
                try:
                    update_organization(org, edit_org)  # Функция для обновления организации в базе данных
                    st.success(f"Организация '{org}' обновлена на '{edit_org}'!")
                    st.rerun()
                except ValueError as e:
                    st.error(str(e))
     else:
        st.warning("Нет организаций, соответствующих вашему запросу.")

 # Удаление организаций
    with st.expander("❌ Удалить организацию", expanded=False):
        orgs = get_organizations()
        if orgs:
            cols = st.columns([3, 1])
            selected_org = cols[0].selectbox(
                "Выберите организацию для удаления", 
                orgs
            )
            if cols[1].button("🗑️ Удалить"):
                delete_organization(selected_org)
                st.success("Организация удалена!")
                st.rerun()
        else:
            st.warning("Нет зарегистрированных организаций")

# --------------------------
# Общие вспомогательные функции
# --------------------------

def save_uploaded_file(uploaded_file):
    if uploaded_file is not None:
        file_ext = uploaded_file.name.split('.')[-1]
        filename = f"{uuid.uuid4()}.{file_ext}"
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        with open(file_path, 'wb') as f:
            f.write(uploaded_file.getbuffer())
        return file_path
    return None

def generate_act(record):
    try:
        doc = Document("template.docx")
        replacements = {
            "{inspection_date}": record.get("inspection_date", ""),
            "{object}": record.get("object", ""),
            "{section}": record.get("section", ""),
            "{organization}": record.get("organization", ""),
            "{violator_name}": record.get("violator_name", ""),
            "{violation_description}": record.get("violation_description", ""),
            "{violation_type}": record.get("violation_type", ""),
            "{violation_category}": record.get("violation_category", ""),
            "{risk_level}": record.get("risk_level", ""),
            "{inspector_name}": record.get("inspector_name", ""),
            "{elimination_date}": record.get("elimination_date", ""),
            "{elimination_status}": record.get("elimination_status", "")
        }
        
        for para in doc.paragraphs:
            for key, value in replacements.items():
                if key in para.text:
                    para.text = para.text.replace(key, str(value))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, str(value))

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"Ошибка генерации акта: {str(e)}")
        return None

# --------------------------
# Запуск приложения
# --------------------------

if __name__ == "__main__":
    main_menu()
    
    if not hasattr(st.session_state, 'module'):
        st.session_state.module = None
        
    if st.session_state.module == "module1":
        module1()
    elif st.session_state.module == "module2":
        module2()
    elif st.session_state.module == "module3":
        module3()
    else:
        st.title("🏠 Главное меню")
        st.write("Выберите модуль в боковом меню 👈")
